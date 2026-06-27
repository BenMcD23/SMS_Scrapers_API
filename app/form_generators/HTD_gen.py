"""ACCTS 7101 Home-to-Duty travel claim filler.

Unlike F1771e (dynamic journey rows), the 7101 is a fixed form — just
placeholder substitution. Handles the template's mixed brace styles:
double `{{ rank }}`, single `{ sn_1 }`, and the malformed `{{ num_j_2 }`
(one closing brace) with a single regex.
"""
import re

from docx import Document

from form_generators.F1771e_gen import _merge_runs

RATE_PER_MILE = 0.25
UPLIFT = 1.07  # 25p/mile, then +7% on top

# {{ key }} | { key } | {{ key }  — one capture, tolerant of the brace count.
_PLACEHOLDER = re.compile(r"\{\{?\s*(\w+)\s*\}\}?")


def compute_htd(distance: float, journeys: list[int]) -> dict:
    """Pure money logic. `journeys` = nights attended per month (≤ 6 entries).

    car_cost = miles × 25p; Total A = car_cost + 7% (applied per return journey);
    each month's amount = journeys × Total A; total claimed = sum of amounts.
    """
    car_cost = round(distance * RATE_PER_MILE, 2)
    total_a = round(distance * RATE_PER_MILE * UPLIFT, 2)
    amounts = [round(j * total_a, 2) for j in journeys]
    return {
        "car_cost": car_cost,
        "total_a": total_a,
        "amounts": amounts,
        "totals": list(amounts),  # claim-for-self only; passenger column blank
        "total_claimed": round(sum(amounts), 2),
    }


def _fill_paragraph(para, context: dict):
    if "{" not in para.text:
        return
    _merge_runs(para)  # reassemble placeholders Word split across runs
    for run in para.runs:
        if "{" in run.text:
            run.text = _PLACEHOLDER.sub(
                lambda m: str(context[m.group(1)]) if m.group(1) in context else m.group(0),
                run.text,
            )


def _fill_tables(tables, context: dict):
    for table in tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:  # merged cells appear once per span
                    continue
                seen.add(id(cell._tc))
                for para in cell.paragraphs:
                    _fill_paragraph(para, context)
                if cell.tables:
                    _fill_tables(cell.tables, context)


def fill_form(template_path: str, output, context: dict):
    """Fill the 7101 template and write to `output` (path or file-like)."""
    doc = Document(template_path)
    for para in doc.paragraphs:
        _fill_paragraph(para, context)
    _fill_tables(doc.tables, context)
    doc.save(output)
