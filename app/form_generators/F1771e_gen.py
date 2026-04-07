"""
fill_f1771e.py
==============
Autofill script for the F1771e Travelling and Subsistence Claim form.

Usage
-----
    from fill_f1771e import fill_form

    personal = {
        "rank":        "Flt Lt",       # see RANKS list for valid values
        "initials":    "J",
        "surname":     "Smith",
        "jpa_number":  "1234567",
        "appointment": "OC 317 Sqn",
        "car_reg":     "AB12 CDE",
    }

    journeys = [
        {
            "date":         "01/04/25",
            "time_depart":  "09:00",
            "time_arrive":  "10:30",
            "from":         "RAF Somewhere, AB1 2CD",
            "to":           "RAF Elsewhere, EF3 4GH",
            "activity":     "Annual Camp — SMS12345",
            "name_rank_pax":"Flt Lt J Smith",
            "hotel_ref":    "",                   # optional
            "misc_expenses":"",                   # optional
            "passengers":   "1",
            "method":       "Car",
            "miles":        "85",
        },
    ]

    fill_form(
        template_path="Accts_Form_F1771e__1_.docx",
        output_path="filled_form.docx",
        personal=personal,
        journeys=journeys,
    )

Column mapping for Part 4 table (13 underlying grid columns):
    C0  — row marker (always empty)
    C1  — Date of journey (DD/MM/YY)
    C2  — Time of Departure
    C3  — Time of Arrival  (gridSpan=2, occupies C3+C4)
    C5  — From
    C6  — To
    C7  — Nature of Activity / Name+Rank / Hotel Ref / Misc  (multi-paragraph)
    C8  — Number of passengers
    C9  — Method
    C10 — Mileage Claimed
    C11 — Accounts Use Only (leave blank)
    C12 — empty
"""

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Valid rank values
# ---------------------------------------------------------------------------
RANKS = [
    "   ", "Sgt", "FS", "WO", "Plt Off", "Fg Off", "Flt Lt",
    "Sqn Ldr", "Wg Cdr", "Gp Capt", "Chaplain", "CI", "CGI",
    "National Chair", "Rgnl Chair", "Wg Chair", "Sqn Chair",
    "Rgnl Treasurer", "Wg Treasurer",
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _merge_runs(para):
    """
    Merge all runs in a paragraph into a single run, preserving the
    formatting of the first run.  Needed because Word sometimes splits
    a placeholder like {{ surname }} across multiple runs.
    """
    runs = para.runs
    if len(runs) <= 1:
        return
    full_text = ''.join(r.text for r in runs)
    # Keep first run's formatting, set its text to the merged string
    runs[0].text = full_text
    # Remove all subsequent runs
    for run in runs[1:]:
        run._r.getparent().remove(run._r)


def _replace_placeholder(doc, placeholder: str, value: str):
    """Replace every occurrence of {{ placeholder }} in all table cells."""
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                if placeholder in cell.text:
                    for para in cell.paragraphs:
                        # Merge runs first in case Word split the placeholder
                        _merge_runs(para)
                        for run in para.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)


def _set_rank_cell(cell, rank: str):
    """
    Set the rank cell. Tries FORMDROPDOWN first; falls back to plain text.
    """
    if rank not in RANKS:
        raise ValueError(f"'{rank}' is not a valid rank. Choose from: {RANKS[1:]}")

    index = RANKS.index(rank)
    tc = cell._tc

    # Try FORMDROPDOWN path
    for fldChar in tc.findall('.//' + qn('w:fldChar')):
        ffData = fldChar.find(qn('w:ffData'))
        if ffData is None:
            continue
        ddList = ffData.find(qn('w:ddList'))
        if ddList is None:
            continue
        result = ddList.find(qn('w:result'))
        if result is None:
            result = OxmlElement('w:result')
            ddList.append(result)
        result.set(qn('w:val'), str(index))
        return

    # Plain-text fallback — write with explicit Arial formatting
    for para in cell.paragraphs[1:]:
        para._p.getparent().remove(para._p)
    para = cell.paragraphs[0]
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r = OxmlElement("w:r")
    r.append(_make_rPr())
    t = OxmlElement("w:t")
    t.text = rank
    r.append(t)
    para._p.append(r)


def _make_rPr():
    """Return a clean <w:rPr>: Arial 10pt, no bold."""
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        fonts.set(qn(attr), 'Arial')
    rPr.append(fonts)
    for tag in ('w:sz', 'w:szCs'):
        el = OxmlElement(tag)
        el.set(qn('w:val'), '20')  # 10pt = 20 half-points
        rPr.append(el)
    return rPr


def _set_tc_text(tc, value: str, extra_lines: list = None):
    """
    Write value into the first paragraph of a <w:tc> using clean Arial
    10pt non-bold formatting. Optionally append extra_lines as additional
    paragraphs (for multi-line cells).
    """
    paras = tc.findall(qn('w:p'))
    if not paras:
        return

    # --- first paragraph ---
    p = paras[0]
    for r in p.findall(qn('w:r')):
        p.remove(r)

    if value:
        r = OxmlElement('w:r')
        r.append(_make_rPr())
        t = OxmlElement('w:t')
        t.text = str(value)
        if value != value.strip():
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)

    # Remove any leftover extra paragraphs from the template
    for extra_p in paras[1:]:
        tc.remove(extra_p)

    # Append extra lines as new paragraphs
    if extra_lines:
        for line in extra_lines:
            if not line:
                continue
            new_p = copy.deepcopy(p)
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            r = OxmlElement('w:r')
            r.append(_make_rPr())
            t = OxmlElement('w:t')
            t.text = str(line)
            r.append(t)
            new_p.append(r)
            tc.append(new_p)


def _bump_bookmark_ids(tr_element, offset: int):
    """Shift all bookmark ids in a cloned row by offset to keep them unique."""
    for bm in tr_element.findall('.//' + qn('w:bookmarkStart')):
        try:
            bm.set(qn('w:id'), str(int(bm.get(qn('w:id'), 0)) + offset))
        except ValueError:
            pass
    for bm in tr_element.findall('.//' + qn('w:bookmarkEnd')):
        try:
            bm.set(qn('w:id'), str(int(bm.get(qn('w:id'), 0)) + offset))
        except ValueError:
            pass


# ---------------------------------------------------------------------------
# TC index → journey dict key  (unique <w:tc> elements in the placeholder row)
# ---------------------------------------------------------------------------
_TC_MAP = {
    1:  'date',
    2:  'time_depart',
    3:  'time_arrive',   # gridSpan=2, still only 1 <w:tc> element
    4:  'from',
    5:  'to',
    # 6 handled separately (multi-line activity cell)
    7:  'passengers',
    8:  'method',
    9:  'miles',
    # 0, 10, 11 — always blank
}


def _build_journey_row(template_tr, journey: dict, bm_offset: int):
    """Clone the placeholder row and fill it with one journey's data."""
    new_tr = copy.deepcopy(template_tr)
    _bump_bookmark_ids(new_tr, bm_offset)

    tcs = new_tr.findall(qn('w:tc'))

    for tc_idx, tc in enumerate(tcs):
        if tc_idx == 6:
            # Multi-line activity cell
            _set_tc_text(
                tc,
                journey.get('activity', ''),
                extra_lines=[
                    journey.get('name_rank_pax', ''),
                    journey.get('hotel_ref', ''),
                    journey.get('misc_expenses', ''),
                ],
            )
        elif tc_idx in _TC_MAP:
            _set_tc_text(tc, journey.get(_TC_MAP[tc_idx], ''))

    return new_tr


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def fill_form(
    template_path: str,
    output_path: str,
    personal: dict,
    journeys: list,
):
    """
    Fill the F1771e form and save to output_path.

    Parameters
    ----------
    template_path : str
        Path to the template .docx. Must have one blank placeholder row
        at the bottom of the Part 4 table (Table 2, Row 3).
    output_path : str
        Destination path for the filled form.
    personal : dict
        Keys: rank, initials, surname, jpa_number, appointment, car_reg
    journeys : list of dict
        Keys: date, time_depart, time_arrive, from, to, activity,
              name_rank_pax, hotel_ref (optional), misc_expenses (optional),
              passengers, method, miles
    """
    doc = Document(template_path)

    # ------------------------------------------------------------------
    # Part 1 — Personal details (Table 0)
    # ------------------------------------------------------------------
    t0 = doc.tables[0]

    rank = personal.get('rank', '')
    if rank:
        _set_rank_cell(t0.rows[1].cells[1], rank)

    _replace_placeholder(doc, '{{ initials }}',    personal.get('initials', ''))
    _replace_placeholder(doc, '{{ surname }}',     personal.get('surname', ''))
    _replace_placeholder(doc, '{{ JPA_num }}',     personal.get('jpa_number', ''))
    _replace_placeholder(doc, '{{ appointment }}', personal.get('appointment', ''))
    _replace_placeholder(doc, '{{ car_reg }}',     personal.get('car_reg', ''))

    # ------------------------------------------------------------------
    # Part 4 — Journey table (Table 2)
    # ------------------------------------------------------------------
    t2 = doc.tables[2]

    # Row 3 is the blank placeholder — insert cloned data rows before it
    # so the blank always stays at the bottom as a visual buffer
    placeholder_tr = t2.rows[3]._tr

    for i, journey in enumerate(journeys):
        new_tr = _build_journey_row(placeholder_tr, journey, bm_offset=(i + 1) * 100)
        placeholder_tr.addprevious(new_tr)

    doc.save(output_path)
    print(f"Saved → {output_path}")
