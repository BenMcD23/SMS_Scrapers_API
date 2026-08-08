"""Placeholder substitution for Word templates authored with ``{{ key }}`` marks.

Shared by every fixed-layout template (the 7101 HTD claim, the NCO appraisal) so
the two fiddly bits live in one place: Word splits a placeholder across runs
whenever you edit it, and merged table cells come back once per column they
span. Templates with a *dynamic* number of rows (F1771e) build their tables in
code instead and don't come through here.
"""

import re

from docx import Document

from form_generators.F1771e_gen import _merge_runs

# {{ key }} | { key } | {{ key }  — one capture, tolerant of the brace count,
# because hand-edited templates end up with all three.
PLACEHOLDER = re.compile(r"\{\{?\s*(\w+)\s*\}\}?")


def fill_paragraph(para, context: dict):
    if "{" not in para.text:
        return
    _merge_runs(para)  # reassemble placeholders Word split across runs
    for run in para.runs:
        if "{" in run.text:
            # python-docx turns "\n" in run text into a <w:br/>, so multi-line
            # values (the appraisal's sections) keep their line breaks.
            run.text = PLACEHOLDER.sub(
                lambda m: str(context[m.group(1)]) if m.group(1) in context else m.group(0),
                run.text,
            )


def fill_tables(tables, context: dict):
    for table in tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:  # merged cells appear once per span
                    continue
                seen.add(id(cell._tc))
                for para in cell.paragraphs:
                    fill_paragraph(para, context)
                if cell.tables:
                    fill_tables(cell.tables, context)


def fill_template(template_path: str, output, context: dict):
    """Fill a template's paragraphs and tables, writing to ``output`` (a path or
    a file-like object)."""
    doc = Document(template_path)
    for para in doc.paragraphs:
        fill_paragraph(para, context)
    fill_tables(doc.tables, context)
    doc.save(output)
