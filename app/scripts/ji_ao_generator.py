from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import io
from datetime import datetime, timedelta


from database.database import SessionLocal
from database.models import Event317


contacts = {
    "McDonald": {
        "email": "ben.mcdonald100@rafac.mod.gov.uk",
        "phone": "07743443608"
    },
    "Stone": {
        "email": "sophie.stone101@rafac.mod.gov.uk",
        "phone": "07735218557"
    },
    "Morris": {
        "email": "gareth.lloyd-morris100@rafac.mod.gov.uk",
        "phone": "07940258406"
    },
    "Doherty": {
        "email": "oc.317@rafac.mod.gov.uk",
        "phone": "07807809776"
    },
    "Gill": {
        "email": "joseph.gill100@rafac.mod.gov.uk",
        "phone": "07543659277"
    },
    "MacGregor": {
        "email": "calum.macgregor100@rafac.mod.gov.uk",
        "phone": "07944026545"
    },
    "Barker": {
        "email": "jonathon.barker100@rafac.mod.gov.uk",
        "phone": "07955063409"
    },
    "Tyrell": {
        "email": "llerytvanessa@gmail.com",
        "phone": "07514586684"
    },
    "N/A": {
        "email": "",
        "phone": ""
    }
}

# Path helpers
def get_template_path(filename):
    return os.path.abspath(
        os.path.join(os.path.dirname(__file__), "..", "word_templates", filename)
    )

def get_signature_path(last_name):
    return os.path.abspath(
        os.path.join(os.path.dirname(__file__), "..", "signatures", f"{last_name}.png")
    )


def replace_text_preserve_format(paragraph, replacements):
    """
    Replace placeholders in a paragraph, preserving formatting and images.
    Handles placeholders split across multiple runs or containing hidden characters.
    """
    for key, value in replacements.items():
        full_text = ''.join(run.text for run in paragraph.runs)
        if key in full_text:
            new_text = full_text.replace(key, str(value))
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

def replace_placeholder_with_signature(paragraph, placeholder, name, email, signature_path, width=Inches(2)):
    """
    Replace a placeholder in a paragraph with a signature image followed by text.
    """
    if placeholder not in paragraph.text:
        return

    # Clear existing text in the paragraph
    paragraph.text = ""

    run = paragraph.add_run()
    try:
        # Add the signature image
        run.add_picture(signature_path, width=width)
    except Exception as e:
        print(f"Could not add signature image for {name}: {e}")
        return
    
    # Add a line break before the text
    paragraph.add_run().add_break()

    # Add text immediately after the image
    paragraph.add_run(f"{name} RAFAC")


def _location_text(event) -> str:
    if not event.location:
        return "N/A"
    first_line = (getattr(event.location, "first_line", "") or "").strip()
    postcode = (getattr(event.location, "postcode", "") or "").strip()
    return f"{first_line}, {postcode}" if first_line and postcode else first_line or postcode or "N/A"


def _date_range_text(event) -> str:
    if not (event.date_from and event.date_to):
        return "N/A"
    if event.date_from.date() == event.date_to.date():
        return event.date_from.strftime("%d/%m/%Y")
    return f"{event.date_from.strftime('%d/%m/%Y')} - {event.date_to.strftime('%d/%m/%Y')}"


def _signature_key(adult_ic: str) -> str:
    """Surname the signature image and contact card are filed under."""
    parts = (adult_ic or "").strip().split()
    return parts[-1] if parts else "N/A"


def _contact(adult_ic: str, key: str) -> str:
    return contacts.get(_signature_key(adult_ic), {}).get(key, "")


def _fmt(dt, pattern: str, fallback: str = "N/A") -> str:
    return dt.strftime(pattern) if dt else fallback


def ji_fields(event) -> dict:
    """Every JI value the template can show, as plain editable text.

    This is both what the UI pre-fills its preview with and what generate_ji
    falls back to for anything the UI doesn't send, so what someone sees on
    screen and what lands in the .docx cannot drift apart.
    """
    cost = event.cost or 0
    at_squadron = bool(event.location) and getattr(event.location, "first_line", "") == "317 Squadron HQ"
    return {
        "title":           event.title or "",
        "date_from_to":    _date_range_text(event),
        "description":     event.description or "",
        "location":        _location_text(event),
        "arrival_time":    _fmt(event.date_from, "%H:%M"),
        "arrival_date":    _fmt(event.date_from, "%d/%m/%Y"),
        "departure_time":  _fmt(event.date_to, "%H:%M"),
        "departure_date":  _fmt(event.date_to, "%d/%m/%Y"),
        "cost": (
            f"Cadets are required to pay £{cost:.2f} to attend this event. "
            "This can be paid via cash/card at squadron or through BACS."
            if cost > 0 else
            "There is no cost for cadets to attend this event."
        ),
        "dress":           event.dress or "",
        "adult_ic":        event.adult_ic or "",
        "adult_ic_email":  _contact(event.adult_ic, "email"),
        "tg_form_req":     "TG 21/23 Forms are not required" if at_squadron else "TG 21/23 Forms are required",
    }


def ao_fields(event) -> dict:
    """Every AO value the template can show. `description` is the free-text
    Activity Description section, which the AO template has no placeholder for —
    it is inserted above the signature only when non-empty."""
    return {
        "todays_date":           datetime.today().strftime("%d %B %Y"),
        "event_ref":             event.reference or "",
        "event_title":           event.title or "",
        "event_location":        _location_text(event),
        "date_from":             _fmt(event.date_from, "%d/%m/%Y", "0"),
        "date_to":               _fmt(event.date_to, "%d/%m/%Y", "0"),
        "course_ic":             f"{event.adult_ic or ''} - {_contact(event.adult_ic, 'email')}",
        "instructor_start_time": _fmt(event.date_from - timedelta(minutes=30) if event.date_from else None, "%H:%M"),
        "cadet_start_time":      _fmt(event.date_from, "%H:%M"),
        "departure_time":        _fmt(event.date_to, "%H:%M"),
        "adult_ic":              event.adult_ic or "",
        "description":           "",
    }


def _apply_signature(doc, adult_ic: str):
    for paragraph in doc.paragraphs:
        if "{{ adult_ic_signature }}" in paragraph.text:
            replace_placeholder_with_signature(
                paragraph,
                "{{ adult_ic_signature }}",
                name=adult_ic,
                email=_contact(adult_ic, "email"),
                signature_path=get_signature_path(_signature_key(adult_ic)),
                width=Inches(2),
            )


def _render(template_name: str, values: dict):
    template_path = get_template_path(template_name)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    doc = Document(template_path)
    replacements = {f"{{{{ {key} }}}}": value for key, value in values.items()}
    for paragraph in doc.paragraphs:
        replace_text_preserve_format(paragraph, replacements)
    return doc


def _save(doc) -> io.BytesIO:
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_ji(event, fields=None):
    """Generate a JI for the selected event.

    `fields` are the edited values from the UI; anything missing falls back to
    what ji_fields computes from the event.
    """
    values = {**ji_fields(event), **(fields or {})}
    doc = _render("ji_template.docx", values)
    _apply_signature(doc, values["adult_ic"])
    return _save(doc)


def generate_ao(event, fields=None):
    """Generate an AO for the selected event. See generate_ji for `fields`."""
    values = {**ao_fields(event), **(fields or {})}
    doc = _render("ao_template.docx", values)

    # The AO template has no free-text placeholder of its own, so a description
    # is inserted as new paragraphs above the signature instead.
    description = (values.get("description") or "").strip()
    if description:
        for paragraph in doc.paragraphs:
            if "{{ adult_ic_signature }}" in paragraph.text:
                paragraph.insert_paragraph_before("Activity Description:").runs[0].bold = True
                paragraph.insert_paragraph_before(description)
                paragraph.insert_paragraph_before("")
                break

    _apply_signature(doc, values["adult_ic"])
    return _save(doc)
